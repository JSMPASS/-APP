"""Word 导出：把一段日期范围内的打卡情况整理成 .docx 文档。

文档结构：标题 -> 统计表 -> 已完成项详情 -> 题目与解析（含错题反思）-> 未完成项列表。
"""
from __future__ import annotations

import os
import shutil
import tempfile
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from habit_checkin.services.export_common import (
    default_filename as _default_filename,
    fmt_duration,
    load_report_data,
    prepare_image,
    report_title,
    result_text,
    weekday_cn,
)
from habit_checkin.ui.richtext import to_plain


def _set_cn_font(run, name="微软雅黑", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _style_document(doc):
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(11)


def _embed_images(doc, tmpdir, db, images):
    for img in images:
        abs_path = db.abs_path(img["file_path"])
        if not os.path.isfile(abs_path):
            continue
        try:
            tmp_img, _, _ = prepare_image(abs_path, tmpdir)
            pic_p = doc.add_paragraph()
            pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pic_p.add_run().add_picture(tmp_img, width=Cm(14))
        except Exception:
            continue


def export_docx(db, start_date, end_date, out_path):
    """导出打卡汇总文档（Word），返回统计 dict。"""
    data = load_report_data(db, start_date, end_date)
    items = data["items"]
    done_items = data["done_items"]
    todo_items = data["todo_items"]
    questions = data["questions"]
    images_map = data["images_map"]
    qimages = data["qimages"]
    by_date = data["by_date"]

    doc = Document()
    _style_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn_font(p.add_run(report_title(start_date, end_date)), size=18, bold=True)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_cn_font(
        sub.add_run("生成时间：{}".format(datetime.now().strftime("%Y-%m-%d %H:%M"))),
        size=10, color=RGBColor(0x80, 0x80, 0x80),
    )

    # 统计表
    total = data["total"]
    rate = data["rate"]
    table = doc.add_table(rows=2, cols=5)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    total_secs = data["total_secs"]
    headers = ["计划项数", "已完成", "完成率", "收录题目", "总学习时长"]
    values = [str(total), str(len(done_items)), "{:.1f}%".format(rate),
              str(len(questions)), fmt_duration(total_secs)]
    for i, (h, v) in enumerate(zip(headers, values)):
        cell = table.cell(0, i)
        _set_cn_font(cell.paragraphs[0].add_run(h), size=11, bold=True)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2 = table.cell(1, i)
        _set_cn_font(cell2.paragraphs[0].add_run(v), size=11)
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    tmpdir = tempfile.mkdtemp(prefix="habit_docx_")
    try:
        # 一、已完成打卡详情
        if done_items:
            h = doc.add_paragraph()
            _set_cn_font(h.add_run("一、已完成打卡详情"), size=14, bold=True)
        for day in sorted(by_date):
            day_p = doc.add_paragraph()
            _set_cn_font(
                day_p.add_run("{}  {}".format(day, weekday_cn(day))),
                size=13, bold=True, color=RGBColor(0x2E, 0x74, 0xB5),
            )
            for it in by_date[day]:
                title = doc.add_paragraph()
                checked = (it["checked_at"] or "")[11:16]
                elapsed = int(it.get("elapsed_seconds") or 0)
                title_suffix = " · 学习时长 {}".format(fmt_duration(elapsed)) if elapsed > 0 else ""
                _set_cn_font(
                    title.add_run("{}（打卡时间 {}{}）".format(it["topic_path"], checked, title_suffix)),
                    size=12, bold=True,
                )
                note = to_plain(it.get("note") or "").strip()
                note_p = doc.add_paragraph()
                _set_cn_font(note_p.add_run("文字总结："), size=11, bold=True)
                if note:
                    _set_cn_font(note_p.add_run(note), size=11)
                else:
                    _set_cn_font(note_p.add_run("（未填写）"), size=11, color=RGBColor(0x80, 0x80, 0x80))
                _embed_images(doc, tmpdir, db, images_map.get(it["id"], []))
                doc.add_paragraph()

        # 二、题目与解析
        if questions:
            hq = doc.add_paragraph()
            _set_cn_font(hq.add_run("二、题目与解析（{} 题）".format(len(questions))), size=14, bold=True)
            for q in questions:
                qt = doc.add_paragraph()
                _set_cn_font(
                    qt.add_run("【{}】{}（{}）".format(q["code"], q["topic_path"], result_text(q))),
                    size=12, bold=True,
                )
                qp = doc.add_paragraph()
                _set_cn_font(qp.add_run("题目内容："), size=11, bold=True)
                question_text = to_plain(q.get("question_text") or "").strip()
                if question_text:
                    _set_cn_font(qp.add_run(question_text), size=11)
                else:
                    _set_cn_font(qp.add_run("（未填写）"), size=11, color=RGBColor(0x80, 0x80, 0x80))
                ap = doc.add_paragraph()
                _set_cn_font(ap.add_run("解析："), size=11, bold=True)
                analysis = to_plain(q.get("analysis") or "").strip()
                if analysis:
                    _set_cn_font(ap.add_run(analysis), size=11)
                else:
                    _set_cn_font(ap.add_run("（未填写）"), size=11, color=RGBColor(0x80, 0x80, 0x80))
                _embed_images(doc, tmpdir, db, qimages.get(q["id"], []))
                filled = bool(
                    to_plain(q.get("self_analysis") or "").strip()
                    or to_plain(q.get("correct_analysis") or "").strip()
                    or to_plain(q.get("reflection") or "").strip()
                )
                if q["result"] == "wrong" or filled:
                    if not filled:
                        rp = doc.add_paragraph()
                        _set_cn_font(rp.add_run("复盘：（未填写）"), size=11, bold=True,
                                     color=RGBColor(0x80, 0x80, 0x80))
                    else:
                        rp = doc.add_paragraph()
                        _set_cn_font(rp.add_run("复盘："), size=11, bold=True)
                        for label, key in (
                            ("自己的做题思路", "self_analysis"),
                            ("正确的做题思路", "correct_analysis"),
                            ("复盘心得", "reflection"),
                        ):
                            val = to_plain(q.get(key) or "").strip()
                            rp2 = doc.add_paragraph()
                            _set_cn_font(rp2.add_run("　{}：".format(label)), size=11, bold=True)
                            _set_cn_font(rp2.add_run(val or "（未填写）"), size=11)
                doc.add_paragraph()
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)

    # 三、未完成项
    if todo_items:
        h2 = doc.add_paragraph()
        _set_cn_font(h2.add_run("三、未完成项（{} 项）".format(len(todo_items))), size=14, bold=True)
        for it in todo_items:
            line = doc.add_paragraph()
            _set_cn_font(
                line.add_run("· {}  {}（未完成）".format(it["plan_date"], it["topic_path"])),
                size=11,
            )

    doc.save(out_path)
    return {"total": total, "done": len(done_items), "rate": rate, "questions": len(questions)}


def default_filename(start_date, end_date):
    return _default_filename(start_date, end_date, "docx")
