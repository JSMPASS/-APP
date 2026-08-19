# -*- coding: utf-8 -*-
"""备考计划文档模板生成与导入同步。

支持格式：
- Markdown (.md / .txt)
- Word (.docx)
- PDF (.pdf，需要 pypdf 用于读取文本)

模板格式示例：
    # 习惯打卡 90 天计划
    - 开始日期：YYYY-MM-DD
    - 结束日期：YYYY-MM-DD
    - 总天数：90

    ## 每日任务
    ### YYYY-MM-DD
    - 09:00 | 政治理论·常识积累 | 辅助 | 小黑课1节 + 要点笔记
    - 09:30 | 判断 | 主 | 花生判断第01讲 + 例题重做
"""
from __future__ import annotations

import re
from datetime import date, timedelta
from pathlib import Path

from habit_checkin.services.study_plan import (
    DEFAULT_START,
    normalize_plan_config,
    resolve_topic,
)

_DATE_RE = re.compile(r"开始(?:日期)?[:：]\s*(\d{4}-\d{2}-\d{2})")
_DAY_RE = re.compile(r"^\s*#+\s*(\d{4}-\d{2}-\d{2})\s*$|^\s*(\d{4}-\d{2}-\d{2})\s*$")
_TASK_RE = re.compile(
    r"^\s*(?:[-*]\s*)?(\d{2}:\d{2})\s*[|｜]\s*([^|｜]+?)\s*[|｜]\s*(主|辅|main|aux)?\s*[|｜]?\s*(.*)$"
)
_HEADING_RE = re.compile(r"^\s*(#{1,6})\s+(.+?)\s*$")
_STAGE_RE = re.compile(
    r"^\s*(?:[-*]\s*)?(?P<name>.+?)（第\s*(?P<day_start>\d+)\s*[-—~]\s*"
    r"(?P<day_end>\d+)\s*天）\s*[:：]?\s*(?P<rest>.*)$"
)
_WEEK_RE = re.compile(r"^\s*(?:[-*]\s*)?第\s*(?P<week>\d+)\s*周\s*[:：]\s*(?P<focus>.*)$")
_CHECKPOINT_RE = re.compile(r"^\s*(?:[-*]\s*)?第\s*(?P<day>\d+)\s*天\s*[:：]\s*(?P<content>.*)$")
_ROUTINE_RE = re.compile(r"^\s*(?:[-*]\s*)?(?P<time>\d{1,2}:\d{2})\s+(?P<desc>.*)$")

_CONFIG_SECTIONS = ("阶段安排", "每周计划", "检查点", "每日作息模板")


# ---------- 读取文档文本 ----------
def read_document_text(path):
    """从 md/docx/pdf 中提取纯文本。"""
    p = Path(path)
    suffix = p.suffix.lower()
    if suffix in (".md", ".txt", ".text"):
        return p.read_text(encoding="utf-8", errors="replace")
    if suffix == ".docx":
        from docx import Document
        doc = Document(str(p))
        parts = [para.text for para in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts)
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
        except ImportError:
            raise RuntimeError("读取 PDF 需要 pypdf，请先安装：python -m pip install pypdf")
        reader = PdfReader(str(p))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    raise ValueError("不支持的文件格式：{}".format(suffix))


# ---------- 模板生成 ----------
def _template_lines(config=None):
    cfg = normalize_plan_config(config)
    end = DEFAULT_START + timedelta(days=cfg["total_days"] - 1)
    lines = [
        "# 习惯打卡计划模板",
        "",
        "- 开始日期：{}".format(DEFAULT_START.isoformat()),
        "- 结束日期：{}".format(end.isoformat()),
        "- 总天数：{}".format(cfg["total_days"]),
        "- 打卡热力图：{} 天".format(cfg["total_days"]),
        "",
        "## 阶段安排",
        "",
    ]
    for s in cfg["stages"]:
        lines.append(_markdown_stage_line(s))
    lines += ["", "## 每周计划", ""]
    for wk, focus in cfg["weeks"]:
        lines.append("- 第 {} 周：{}".format(wk, focus))
    lines += ["", "## 检查点", ""]
    for cp_day, content in cfg["checkpoints"]:
        lines.append("- 第 {} 天：{}".format(cp_day, content))
    lines += ["", "## 每日作息模板", ""]
    for tm, desc in cfg["daily_routine"]:
        lines.append("- {} {}".format(tm, desc))
    lines += [
        "",
        "## 每日任务",
        "",
        "> 每行格式：`时间 | 任务/知识点 | 主/辅 | 备注`",
        "> 可自行增删日期和任务，修改后重新上传即可同步到 App。",
        "",
        "### {}".format(DEFAULT_START.isoformat()),
        "- 09:00 | 政治理论·常识积累 | 辅助 | 小黑课1节 + 要点笔记",
        "- 09:30 | 判断 | 主 | 花生判断第01讲 + 例题重做",
        "- 11:00 | 资料 | 主 | 20题限时25分钟 + 公式默写",
        "- 13:30 | 归纳概括 | 主 | 袁东·归纳概括方法课 + 归纳概括1题",
        "- 16:00 | 错题复盘 | 辅助 | 错题ABCD分类 + 当日重点整理 + 方法本整理",
        "- 19:00 | 新闻联播 | 辅助 | 收听新闻联播 + 记录时政要点",
        "- 19:30 | 创作 | 辅助 | ComfyUI / 软件编程创作（自由创作 2 小时）",
        "",
        "### {}".format((DEFAULT_START + timedelta(days=1)).isoformat()),
        "- 09:00 | 政治理论·常识积累 | 辅助 | 背诵清单挖空版 + 睡前复习",
        "- 09:30 | 判断 | 主 | 花生判断第02讲 + 例题重做",
        "- 11:00 | 言语 | 主 | 20题限时20分钟 + 成语积累10个",
        "- 13:30 | 归纳概括 | 主 | 袁东·归纳概括方法课 + 归纳概括1题",
        "- 16:00 | 错题复盘 | 辅助 | 错题ABCD分类 + 当日重点整理 + 方法本整理",
        "- 19:00 | 新闻联播 | 辅助 | 收听新闻联播 + 记录时政要点",
        "- 19:30 | 创作 | 辅助 | ComfyUI / 软件编程创作（自由创作 2 小时）",
        "",
        "> 后续日期请按此格式继续添加。",
        "",
    ]
    return lines


def export_markdown_template(path, config=None):
    Path(path).write_text("\n".join(_template_lines(config)), encoding="utf-8")
    return str(path)


def export_docx_template(path, config=None):
    from docx import Document
    from docx.shared import Pt

    cfg = normalize_plan_config(config)
    end = DEFAULT_START + timedelta(days=cfg["total_days"] - 1)
    doc = Document()
    doc.add_heading("习惯打卡计划模板", level=0)
    doc.add_paragraph("开始日期：{}".format(DEFAULT_START.isoformat()))
    doc.add_paragraph("结束日期：{}".format(end.isoformat()))
    doc.add_paragraph("总天数：{}".format(cfg["total_days"]))
    doc.add_heading("阶段安排", level=1)
    for s in cfg["stages"]:
        doc.add_paragraph(_markdown_stage_line(s))
    doc.add_heading("每周计划", level=1)
    for wk, focus in cfg["weeks"]:
        doc.add_paragraph("- 第 {} 周：{}".format(wk, focus))
    doc.add_heading("检查点", level=1)
    for cp_day, content in cfg["checkpoints"]:
        doc.add_paragraph("- 第 {} 天：{}".format(cp_day, content))
    doc.add_heading("每日作息模板", level=1)
    for tm, desc in cfg["daily_routine"]:
        doc.add_paragraph("- {} {}".format(tm, desc))
    doc.add_heading("每日任务", level=1)
    doc.add_paragraph("每行格式：时间 | 任务/知识点 | 主/辅 | 备注")
    table = doc.add_table(rows=8, cols=4)
    table.style = "Light Grid Accent 1"
    headers = ["时间", "任务/知识点", "主/辅", "备注"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    sample = [
        ("09:00", "政治理论·常识积累", "辅助", "小黑课1节 + 要点笔记"),
        ("09:30", "判断", "主", "花生判断第01讲 + 例题重做"),
        ("11:00", "资料", "主", "20题限时25分钟 + 公式默写"),
        ("13:30", "归纳概括", "主", "袁东·归纳概括方法课 + 归纳概括1题"),
        ("16:00", "错题复盘", "辅助", "错题ABCD分类 + 当日重点整理 + 方法本整理"),
        ("19:00", "新闻联播", "辅助", "收听新闻联播 + 记录时政要点"),
        ("19:30", "创作", "辅助", "ComfyUI / 软件编程创作（自由创作 2 小时）"),
    ]
    for r, row in enumerate(sample, start=1):
        for c, val in enumerate(row):
            table.cell(r, c).text = val
    doc.add_paragraph("后续日期请按此格式继续添加。")
    doc.save(str(path))
    return str(path)


def export_pdf_template(path, config=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    from reportlab.lib import colors

    try:
        pdfmetrics.getFont("STSong-Light")
    except Exception:
        pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    font = "STSong-Light"
    styles = {
        "title": ParagraphStyle("title", fontName=font, fontSize=18, leading=24, alignment=1),
        "body": ParagraphStyle("body", fontName=font, fontSize=10.5, leading=16),
        "h": ParagraphStyle("h", fontName=font, fontSize=14, leading=20, spaceBefore=10),
    }
    cfg = normalize_plan_config(config)
    end = DEFAULT_START + timedelta(days=cfg["total_days"] - 1)
    doc = SimpleDocTemplate(str(path), pagesize=A4,
                            leftMargin=2 * cm, rightMargin=2 * cm,
                            topMargin=1.8 * cm, bottomMargin=1.8 * cm)
    story = [
        Paragraph("习惯打卡计划模板", styles["title"]),
        Spacer(1, 12),
        Paragraph("开始日期：{}".format(DEFAULT_START.isoformat()), styles["body"]),
        Paragraph("结束日期：{}".format(end.isoformat()), styles["body"]),
        Paragraph("总天数：{}".format(cfg["total_days"]), styles["body"]),
        Spacer(1, 8),
        Paragraph("阶段安排", styles["h"]),
    ]
    for s in cfg["stages"]:
        story.append(Paragraph(_markdown_stage_line(s), styles["body"]))
    story.append(Paragraph("每周计划", styles["h"]))
    for wk, focus in cfg["weeks"]:
        story.append(Paragraph("- 第 {} 周：{}".format(wk, focus), styles["body"]))
    story.append(Paragraph("检查点", styles["h"]))
    for cp_day, content in cfg["checkpoints"]:
        story.append(Paragraph("- 第 {} 天：{}".format(cp_day, content), styles["body"]))
    story.append(Paragraph("每日作息模板", styles["h"]))
    for tm, desc in cfg["daily_routine"]:
        story.append(Paragraph("- {} {}".format(tm, desc), styles["body"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph("每日任务", styles["h"]))
    story.append(Paragraph("每行格式：时间 | 任务/知识点 | 主/辅 | 备注", styles["body"]))
    story.append(Spacer(1, 8))
    data = [
        ["时间", "任务/知识点", "主/辅", "备注"],
        ["09:00", "政治理论·常识积累", "辅助", "小黑课1节 + 要点笔记"],
        ["09:30", "判断", "主", "花生判断第01讲 + 例题重做"],
        ["11:00", "资料", "主", "20题限时25分钟 + 公式默写"],
        ["13:30", "归纳概括", "主", "袁东·归纳概括方法课 + 归纳概括1题"],
        ["16:00", "错题复盘", "辅助", "错题ABCD分类 + 当日重点整理 + 方法本整理"],
        ["19:00", "新闻联播", "辅助", "收听新闻联播 + 记录时政要点"],
        ["19:30", "创作", "辅助", "ComfyUI / 软件编程创作（自由创作 2 小时）"],
    ]
    table = Table(data, colWidths=[2.2 * cm, 5.5 * cm, 1.8 * cm, 5.5 * cm])
    table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, -1), font),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#EFF3F9")),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D8E0EA")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph("后续日期请按此格式继续添加。", styles["body"]))
    doc.build(story)
    return str(path)


def export_template(path, fmt, config=None):
    fmt = (fmt or "").lower()
    if fmt in ("md", "markdown", "txt"):
        return export_markdown_template(path, config=config)
    if fmt == "docx":
        return export_docx_template(path, config=config)
    if fmt == "pdf":
        return export_pdf_template(path, config=config)
    raise ValueError("不支持的模板格式：{}".format(fmt))


# ---------- 配置段落回写 ----------
def _markdown_stage_line(s):
    return "- {}（第 {} - {} 天）：{}；申论：{}；退出标准：{}".format(
        s["name"], s["day_start"], s["day_end"],
        s.get("xingce", ""), s.get("shenlun", ""), s.get("exit", ""),
    )


def _section_item_lines(heading, cfg):
    if heading == "阶段安排":
        return [_markdown_stage_line(s) for s in cfg["stages"]]
    if heading == "每周计划":
        return ["- 第 {} 周：{}".format(wk, focus) for wk, focus in cfg["weeks"]]
    if heading == "检查点":
        return ["- 第 {} 天：{}".format(cp_day, content) for cp_day, content in cfg["checkpoints"]]
    if heading == "每日作息模板":
        return ["- {} {}".format(tm, desc) for tm, desc in cfg["daily_routine"]]
    return []


def _heading_name(line):
    m = _HEADING_RE.match(line)
    return m.group(2).strip() if m else None


def _item_key(heading, line):
    if heading == "每周计划":
        m = _WEEK_RE.match(line.strip())
        return int(m.group("week")) if m else None
    if heading == "检查点":
        m = _CHECKPOINT_RE.match(line.strip())
        return int(m.group("day")) if m else None
    if heading == "每日作息模板":
        m = _ROUTINE_RE.match(line.strip())
        return m.group("time") if m else None
    return None


def _sync_section_items(body, heading, cfg):
    """按结构化行回写配置，保留段落里的自定义备注与空行。"""
    if heading == "阶段安排":
        new_items = _section_item_lines(heading, cfg)
        out = []
        idx = 0
        for line in body:
            if _STAGE_RE.match(line.strip()):
                if idx < len(new_items):
                    out.append(new_items[idx])
                    idx += 1
                continue
            out.append(line)
        out.extend(new_items[idx:])
        return out

    keyed = []
    if heading == "每周计划":
        keyed = [(int(wk), "- 第 {} 周：{}".format(wk, focus)) for wk, focus in cfg["weeks"]]
    elif heading == "检查点":
        keyed = [(int(cp_day), "- 第 {} 天：{}".format(cp_day, content))
                 for cp_day, content in cfg["checkpoints"]]
    elif heading == "每日作息模板":
        keyed = [(tm, "- {} {}".format(tm, desc)) for tm, desc in cfg["daily_routine"]]
    wanted = {key: line for key, line in keyed}
    seen = set()
    out = []
    for line in body:
        key = _item_key(heading, line)
        if key is not None:
            if key in wanted and key not in seen:
                out.append(wanted[key])
                seen.add(key)
            continue
        out.append(line)
    out.extend(line for key, line in keyed if key not in seen)
    return out


def update_markdown_config_sections(path, config):
    """把计划配置回写到 Markdown/文本文件的四个配置段落。

    「每日任务」段落及其中的日期、任务、备注保持不变；配置段落里的
    自定义备注行也会保留，只替换/增删结构化条目。
    """
    cfg = normalize_plan_config(config)
    p = Path(path)
    text = p.read_text(encoding="utf-8", errors="replace")
    newline = "\r\n" if "\r\n" in text else "\n"
    lines = text.splitlines()

    for heading in _CONFIG_SECTIONS:
        start = next((i for i, line in enumerate(lines) if _heading_name(line) == heading), None)
        if start is None:
            continue
        j = start + 1
        while j < len(lines) and _heading_name(lines[j]) is None:
            j += 1
        body = _sync_section_items(lines[start + 1:j], heading, cfg)
        lines = lines[:start + 1] + body + lines[j:]

    missing = [h for h in _CONFIG_SECTIONS
               if not any(_heading_name(line) == h for line in lines)]
    if missing:
        block = []
        for heading in missing:
            block += ["## " + heading, ""] + _section_item_lines(heading, cfg) + [""]
        daily_idx = next((i for i, line in enumerate(lines)
                          if _heading_name(line) == "每日任务"), None)
        if daily_idx is None:
            lines += block
        else:
            lines = lines[:daily_idx] + block + lines[daily_idx:]

    content = newline.join(lines)
    if not content.endswith(newline):
        content += newline
    p.write_text(content, encoding="utf-8")
    return str(p)


def markdown_source_path(db):
    """返回可回写的计划来源文件；不是 Markdown/文本或已不存在时返回 None。"""
    rel = db.get_setting("plan_source_file", "") if hasattr(db, "get_setting") else ""
    if not rel:
        return None
    path = Path(db.abs_path(rel))
    if path.suffix.lower() not in (".md", ".txt", ".text"):
        return None
    return str(path) if path.is_file() else None


# ---------- 导入同步 ----------
def _parse_date(text):
    m = _DATE_RE.search(text)
    if m:
        return date.fromisoformat(m.group(1))
    # 退而求其次：取文本中第一个完整日期
    for mm in re.finditer(r"\b(\d{4}-\d{2}-\d{2})\b", text):
        try:
            return date.fromisoformat(mm.group(1))
        except ValueError:
            continue
    return None


def _parse_tasks(text):
    """返回 {date_str: [(time, label, task_type, note), ...]}"""
    tasks = {}
    current = None
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        dm = _DAY_RE.match(line)
        if dm:
            d = dm.group(1) or dm.group(2)
            try:
                date.fromisoformat(d)
                current = d
                tasks.setdefault(current, [])
                continue
            except ValueError:
                pass
        tm = _TASK_RE.match(line)
        if tm and current:
            time_str = tm.group(1)
            label = tm.group(2).strip()
            task_type = (tm.group(3) or "主").strip().lower()
            if task_type in ("辅", "aux"):
                task_type = "aux"
            else:
                task_type = "main"
            note = (tm.group(4) or "").strip()
            tasks[current].append((time_str, label, task_type, note))
    return {k: v for k, v in tasks.items() if v}


def import_plan_document(db, path, overwrite=True):
    """把模板文档同步到数据库：更新开始日期并重建文档中列出的每日计划。"""
    text = read_document_text(path)
    start = _parse_date(text)
    if start is None:
        raise ValueError("文档中没有找到有效的开始日期（例如：开始日期：YYYY-MM-DD）")
    tasks = _parse_tasks(text)
    if not tasks:
        # 只有开始日期时，至少更新 plan_start_date，不重建计划
        db.set_setting("plan_start_date", start.isoformat())
        db.set_setting("plan_source_file", db.rel_path(path))
        return {"start": start.isoformat(), "days": 0, "items": 0, "updated_start_only": True}

    created_days = created_items = 0
    for day_str, rows in tasks.items():
        try:
            d = date.fromisoformat(day_str)
        except ValueError:
            continue
        day_no = (d - start).days + 1
        if day_no < 1:
            continue
        existing = db.get_plan(day_str)
        if existing:
            if not overwrite:
                continue
            db.delete_plan(existing["id"])
        plan_id = db.create_plan(day_str, title="第 {} 天 · 导入计划".format(day_no))
        for time_str, label, task_type, note in rows:
            try:
                topic_id = resolve_topic(db, label, day_no)
            except ValueError:
                parts = [p.strip() for p in label.split("/") if p.strip()]
                if not parts:
                    parts = [label]
                topic_id = db.ensure_topic_by_path(tuple(parts))
            db.add_plan_item(plan_id, topic_id, time_str, task_type, note)
            created_items += 1
        db.conn.commit()
        created_days += 1
    db.set_setting("plan_start_date", start.isoformat())
    db.set_setting("plan_source_file", db.rel_path(path))
    return {"start": start.isoformat(), "days": created_days, "items": created_items, "updated_start_only": False}
