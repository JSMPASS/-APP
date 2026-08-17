"""导出单元测试：Word 含题目/解析/反思，PDF 中文生成。"""
from __future__ import annotations

import tempfile
import unittest
from datetime import date, timedelta
from pathlib import Path

import sys
sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "src"))

from habit_checkin.db import Database
from habit_checkin.services.export_docx import default_filename, export_docx
from habit_checkin.services.export_pdf import default_filename_pdf, export_pdf
from habit_checkin.services.export_common import (
    prepare_image,
    report_title,
    weekday_cn,
)


class TestExport(unittest.TestCase):
    def setUp(self):
        self.tmp = tempfile.TemporaryDirectory()
        self.root = Path(self.tmp.name)
        self.db = Database(self.root / "app.db", self.root / "images", self.root)
        from PIL import Image
        img1 = self.root / "a.png"
        img2 = self.root / "b.jpg"
        Image.new("RGB", (800, 600), (0, 120, 200)).save(img1)
        Image.new("RGB", (300, 500), (200, 80, 40)).save(img2)
        self.img1_rel = self.db.store_image_from_path(str(img1))
        self.img2_rel = self.db.store_image_from_path(str(img2))
        self.today = date.today().isoformat()
        self.tomorrow = (date.today() + timedelta(days=1)).isoformat()

    def tearDown(self):
        self.db.close()
        self.tmp.cleanup()

    def _leaf(self, name):
        return [t for t in self.db.list_topics() if t["name"] == name][0]

    def _setup_data(self):
        pid = self.db.create_plan(self.today, "今日")
        iid = self.db.add_plan_item(pid, self._leaf("逻辑填空（选词填空）")["id"], "19:30")
        self.db.update_checkin(iid, "完成 30 道逻辑填空，正确率 85%。", done=True)
        self.db.add_image(iid, self.img1_rel, sort_order=0)
        self.db.add_image(iid, self.img2_rel, sort_order=1)
        iid2 = self.db.add_plan_item(pid, self._leaf("大作文")["id"])
        # 错题 + 反思
        qid = self.db.add_question(topic_id=self._leaf("单一指标")["id"],
                                   question_text="2023 年增速为 8%，求增量。",
                                   analysis="增量 = 现期 * 增速 / (1+增速)",
                                   result="wrong", result_reason="计算粗心")
        self.db.sync_question_images(qid, [], [str(self.root / "a.png")])
        self.db.update_question(qid, self_analysis="我的思路绕远", correct_analysis="直接套公式",
                                reflection="记住增量公式")
        # 正确题
        self.db.add_question(question_text="判断推理一题", result="correct", result_reason="完全理解")
        return iid, iid2

    def test_export_single_day(self):
        self._setup_data()
        out = self.root / "out.docx"
        stats = export_docx(self.db, self.today, self.today, str(out))
        self.assertTrue(out.is_file())
        self.assertEqual(stats["total"], 2)
        self.assertEqual(stats["done"], 1)
        self.assertEqual(stats["rate"], 50.0)
        self.assertEqual(stats["questions"], 2)
        from docx import Document
        doc = Document(str(out))
        self.assertGreater(len(doc.paragraphs), 5)
        self.assertGreaterEqual(len(doc.tables), 1)
        self.assertGreaterEqual(len(doc.inline_shapes), 3)  # 2 打卡图 + 1 题目图
        texts = "\n".join(p.text for p in doc.paragraphs)
        for key in ["打卡情况汇总", "题目与解析", "Q0001", "Q0002", "复盘",
                    "计算粗心", "我的思路绕远", "正确的做题思路", "未完成", "大作文"]:
            self.assertIn(key, texts)
        self.assertEqual(default_filename(self.today, self.today), "打卡情况_{}.docx".format(self.today))

    def test_export_range_and_pdf(self):
        self._setup_data()
        pid2 = self.db.create_plan(self.tomorrow, "")
        self.db.add_plan_item(pid2, self._leaf("数字推理")["id"], "08:00")
        out = self.root / "range.docx"
        stats = export_docx(self.db, self.today, self.tomorrow, str(out))
        self.assertTrue(out.is_file())
        self.assertEqual(stats["total"], 3)
        self.assertEqual(default_filename(self.today, self.tomorrow), "打卡情况_{}_{}.docx".format(self.today, self.tomorrow))
        # PDF
        out_pdf = self.root / "range.pdf"
        stats_pdf = export_pdf(self.db, self.today, self.tomorrow, str(out_pdf))
        self.assertTrue(out_pdf.is_file())
        self.assertEqual(stats_pdf["questions"], 2)
        from pypdf import PdfReader
        reader = PdfReader(str(out_pdf))
        self.assertGreaterEqual(len(reader.pages), 1)
        first = reader.pages[0].extract_text() or ""
        self.assertIn("打卡情况汇总", first)
        self.assertEqual(default_filename_pdf(self.today, self.tomorrow), "打卡情况_{}_{}.pdf".format(self.today, self.tomorrow))

    def test_export_common_helpers(self):
        from PIL import Image
        src = self.root / "big.png"
        Image.new("RGB", (2400, 1200), (10, 20, 30)).save(src)
        with tempfile.TemporaryDirectory() as td:
            p, w, h = prepare_image(str(src), td)
            self.assertTrue(Path(p).is_file())
            self.assertLessEqual(max(w, h), 1600)
        self.assertEqual(report_title(self.today, self.today), "打卡情况汇总（{}）".format(self.today))
        self.assertEqual(
            report_title(self.today, self.tomorrow),
            "打卡情况汇总（{} 至 {}）".format(self.today, self.tomorrow),
        )
        d = date.fromisoformat(self.today)
        self.assertEqual(weekday_cn(self.today), "星期" + "一二三四五六日"[d.weekday()])


if __name__ == "__main__":
    unittest.main()
